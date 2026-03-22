from __future__ import annotations

import argparse
import re
import shutil
import tempfile
from copy import deepcopy
from pathlib import Path

from lxml import etree
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from docxcompose.composer import Composer
from docx.oxml.ns import qn

BODY_FIRST_LINE_INDENT_PT = 24
SECTION_HEADING_SPACE_PT = 15.6
BIB_HANGING_INDENT_PT = 21


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description="Assemble a thesis DOCX using a senior's DOCX as the style/template base."
    )
    parser.add_argument("--template-docx", required=True)
    parser.add_argument("--main-docx", required=True)
    parser.add_argument("--backmatter-docx", required=True)
    parser.add_argument("--frontpages-tex", required=True)
    parser.add_argument("--output-docx", required=True)
    return parser.parse_args()


def read_frontpage_metadata(path: Path) -> dict[str, str]:
    text = path.read_text(encoding="utf-8")
    keys = [
        "title",
        "chinesetitle",
        "author",
        "advisor",
        "advisorsec",
        "degree",
        "degreetype",
        "major",
        "institute",
        "research",
        "authorno",
        "submissiondate",
        "oraldefencedate",
        "degreedate",
        "chairman",
        "englishtitle",
        "englishauthor",
        "englishadvisor",
        "englishdegree",
        "englishdegreetype",
        "englishmajor",
        "englishinstitute",
        "englishdate",
    ]
    metadata: dict[str, str] = {}
    for key in keys:
        match = re.search(r"\\" + re.escape(key) + r"\{([^}]*)\}", text)
        if match:
            metadata[key] = match.group(1).strip()
    return metadata


def normalize_tex_date(text: str) -> str:
    return text.replace("~", "").strip()


def element_text(element) -> str:
    return "".join(node.text for node in element.iter() if getattr(node, "text", None))


def replace_text_in_paragraph(paragraph, replacements: list[tuple[str, str]]) -> None:
    text = paragraph.text
    updated = text
    for old, new in replacements:
        updated = updated.replace(old, new)
    if updated != text:
        paragraph.text = updated


def replace_text_in_doc(doc: Document, replacements: list[tuple[str, str]]) -> None:
    for paragraph in doc.paragraphs:
        replace_text_in_paragraph(paragraph, replacements)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_text_in_paragraph(paragraph, replacements)


def trim_doc_from_text(doc: Document, marker_text: str) -> None:
    marker_normalized = marker_text.replace("\u3000", "").replace(" ", "").strip()
    body = doc.element.body
    removing = False
    for child in list(body):
        tag = child.tag.rsplit("}", 1)[-1]
        if tag == "sectPr":
            continue
        child_text = element_text(child).replace("\u3000", "").replace(" ", "").strip()
        if not removing and marker_normalized and marker_normalized in child_text:
            removing = True
        if removing:
            body.remove(child)


def remove_body_child(child) -> None:
    parent = child.getparent()
    if parent is not None:
        parent.remove(child)


def normalize_section_text(text: str) -> str:
    return text.replace("\u3000", "").replace(" ", "").strip()


def remove_unwanted_front_sections(doc: Document) -> None:
    start_keywords = ("插图索引", "表格索引", "ListofFigures", "ListofTables")
    stop_keywords = ("摘要", "Abstract")
    figure_table_toc_codes = (
        'TOC \\h \\z \\c "Figure"',
        'TOC \\h \\z \\c Figure',
        'TOC \\h \\z \\c "Table"',
        'TOC \\h \\z \\c Table',
    )

    body = doc.element.body
    removing = False
    for child in list(body):
        tag = child.tag.rsplit("}", 1)[-1]
        if tag == "sectPr":
            continue
        child_text = normalize_section_text(element_text(child))
        child_xml = etree.tostring(child, encoding="unicode")
        has_unwanted_title = any(keyword in child_text for keyword in start_keywords)
        has_unwanted_toc = any(code in child_xml for code in figure_table_toc_codes)
        if has_unwanted_title or has_unwanted_toc:
            removing = True
        if removing:
            if child_text and any(keyword in child_text for keyword in stop_keywords):
                removing = False
                continue
            remove_body_child(child)


def set_run_font(run, east_asia: str, ascii_name: str, size_pt: float, bold: bool) -> None:
    run.bold = bold
    run.font.size = Pt(size_pt)
    run.font.name = ascii_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)


def format_center_heading(paragraph, text: str) -> None:
    paragraph.text = text
    paragraph.style = paragraph.part.document.styles["Normal"]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fmt = paragraph.paragraph_format
    fmt.page_break_before = True
    fmt.first_line_indent = Pt(0)
    fmt.space_before = Pt(SECTION_HEADING_SPACE_PT)
    fmt.space_after = Pt(SECTION_HEADING_SPACE_PT)
    for run in paragraph.runs:
        set_run_font(run, "宋体", "Times New Roman", 16, True)


def renumber_heading(text: str, levels: int) -> str:
    clean = text.replace("\t", " ").strip()
    if levels == 2:
        match = re.match(r"^(\d+)\.(\d+)\s+(.+)$", clean)
        if match:
            chap = int(match.group(1)) - 2
            return f"{chap}.{match.group(2)}\u3000{match.group(3)}"
    if levels == 3:
        match = re.match(r"^(\d+)\.(\d+)\.(\d+)\s+(.+)$", clean)
        if match:
            chap = int(match.group(1)) - 2
            return f"{chap}.{match.group(2)}.{match.group(3)}\u3000{match.group(4)}"
    return text


def style_bibliography_paragraph(paragraph) -> None:
    try:
        paragraph.style = paragraph.part.document.styles["List Paragraph"]
    except KeyError:
        pass
    fmt = paragraph.paragraph_format
    fmt.left_indent = Pt(0)
    fmt.first_line_indent = Pt(-BIB_HANGING_INDENT_PT)
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)


def style_section_heading(paragraph, *, page_break_before: bool = False) -> None:
    fmt = paragraph.paragraph_format
    fmt.first_line_indent = Pt(0)
    fmt.space_before = Pt(SECTION_HEADING_SPACE_PT)
    fmt.space_after = Pt(SECTION_HEADING_SPACE_PT)
    fmt.page_break_before = page_break_before


def style_body_paragraph(paragraph, *, indent_pt: float = BODY_FIRST_LINE_INDENT_PT) -> None:
    fmt = paragraph.paragraph_format
    fmt.left_indent = Pt(0)
    fmt.first_line_indent = Pt(indent_pt)
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)


def style_keyword_paragraph(paragraph) -> None:
    fmt = paragraph.paragraph_format
    fmt.left_indent = Pt(0)
    fmt.first_line_indent = Pt(0)
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)


def is_keyword_line(text: str) -> bool:
    return bool(re.match(r"^(关键词|Key words)[:：]", text))


def is_figure_or_table_caption(text: str) -> bool:
    return bool(
        re.match(r"^(图\d+(\.\d+)?|Fig\.)", text)
        or re.match(r"^(表\d+(\.\d+)?|Table)", text)
    )


def looks_like_body_text(text: str) -> bool:
    compact = re.sub(r"\s+", "", text)
    if len(compact) < 20:
        return False
    if is_keyword_line(text):
        return False
    if is_figure_or_table_caption(text):
        return False
    if re.match(r"^\[\d+\]", text):
        return False
    return True


def strip_reference_url(text: str) -> str:
    cleaned = re.sub(r"\s*(https?://\S+|www\.\S+)\s*", " ", text)
    cleaned = re.sub(r"\.\s*,", ".", cleaned)
    cleaned = re.sub(r"\s+,", ",", cleaned)
    cleaned = re.sub(r"\s+\.", ".", cleaned)
    cleaned = re.sub(r"\s{2,}", " ", cleaned)
    return cleaned.strip()


def transform_main_doc(doc: Document) -> None:
    first_bib_para = None
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        if paragraph.style.name == "Heading 1":
            if re.match(r"^1\s+摘要$", text.replace("\t", " ")):
                format_center_heading(paragraph, "摘　要")
                continue
            if re.match(r"^2\s+Abstract$", text.replace("\t", " ")):
                format_center_heading(paragraph, "Abstract")
                continue
            match = re.match(r"^(\d+)\s+(.+)$", text.replace("\t", " "))
            if match:
                chapter_no = int(match.group(1)) - 2
                title = match.group(2)
                paragraph.text = f"第{chapter_no}章\u3000{title}"
                style_section_heading(paragraph, page_break_before=True)
                continue
        if paragraph.style.name == "Heading 2":
            paragraph.text = renumber_heading(text, 2)
        elif paragraph.style.name == "Heading 3":
            paragraph.text = renumber_heading(text, 3)

        if re.match(r"^\[\d+\]", text):
            cleaned_ref = strip_reference_url(text)
            if cleaned_ref != text:
                paragraph.text = cleaned_ref
                text = cleaned_ref
            if first_bib_para is None:
                first_bib_para = paragraph
            style_bibliography_paragraph(paragraph)

        if paragraph.style.name == "Normal":
            if is_keyword_line(text):
                style_keyword_paragraph(paragraph)
            elif re.match(r"^(图\d+(\.\d+)?|Fig\.)", text):
                try:
                    paragraph.style = doc.styles["图表标题"]
                except KeyError:
                    pass
            elif re.match(r"^(表\d+(\.\d+)?|Table)", text):
                try:
                    paragraph.style = doc.styles["图表标题"]
                except KeyError:
                    pass
            elif looks_like_body_text(text):
                style_body_paragraph(paragraph)

    if first_bib_para is not None:
        heading = first_bib_para.insert_paragraph_before("参考文献", style="Heading 1")
        style_section_heading(heading, page_break_before=True)


def transform_backmatter_doc(doc: Document) -> None:
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text == "[plain]":
            paragraph.text = ""

    heading_count = 0
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        if paragraph.style.name == "Heading 1":
            heading_count += 1
            if heading_count == 1:
                paragraph.text = "致　谢"
                style_section_heading(paragraph, page_break_before=True)
            elif heading_count == 2:
                paragraph.text = "攻读学位期间发表的学术成果"
                style_section_heading(paragraph, page_break_before=True)
        elif paragraph.style.name == "Heading 2" and text == "学术论文:":
            paragraph.style = doc.styles["Normal"]
            style_body_paragraph(paragraph)
            for run in paragraph.runs:
                run.bold = True
        else:
            paragraph.style = doc.styles["Normal"]
            if looks_like_body_text(text) or len(re.sub(r"\s+", "", text)) >= 8:
                style_body_paragraph(paragraph)


def save_temp_doc(doc: Document, path: Path) -> Path:
    doc.save(path)
    return path


def main() -> int:
    args = parse_args()
    template_docx = Path(args.template_docx).resolve()
    main_docx = Path(args.main_docx).resolve()
    backmatter_docx = Path(args.backmatter_docx).resolve()
    frontpages_tex = Path(args.frontpages_tex).resolve()
    output_docx = Path(args.output_docx).resolve()
    output_docx.parent.mkdir(parents=True, exist_ok=True)

    metadata = read_frontpage_metadata(frontpages_tex)
    chinese_date = normalize_tex_date(metadata.get("submissiondate", ""))
    degree_date = normalize_tex_date(metadata.get("degreedate", chinese_date))
    english_date = metadata.get("englishdate", "")
    author = metadata.get("author", "")
    advisor = metadata.get("advisor", "")
    english_author = metadata.get("englishauthor", author)
    english_advisor = metadata.get("englishadvisor", advisor)
    chinese_title = metadata.get("chinesetitle") or metadata.get("title", "")
    english_title = metadata.get("englishtitle", "")
    major = metadata.get("major", "")
    english_major = metadata.get("englishmajor", major)

    replacements = [
        ("面向视觉定位的多模态大模型结构优化与方法研究", chinese_title),
        ("Research on structure optimization and method of multimodal large model for visual grounding", english_title),
        ("鄢锦豪", author),
        ("Yan Jinhao", english_author),
        ("陈东岳  教授", advisor),
        ("陈东岳 教授", advisor),
        ("Professor Chen Dongyue", english_advisor),
        ("Control Engineering", english_major),
        ("控制工程", major),
        ("2270902", metadata.get("authorno", "")),
        ("2025 年 6 月", chinese_date.replace("年", " 年 ").replace("月", " 月").replace("  ", " ").strip()),
        ("2025年6月10日", chinese_date),
        ("2025年6月", chinese_date),
        ("2025年7月", degree_date),
        ("June 2025", english_date),
    ]
    replacements = [(old, new) for old, new in replacements if old and new]
    replacements.sort(key=lambda item: len(item[0]), reverse=True)

    with tempfile.TemporaryDirectory(prefix="thesis-word-build-") as tmp_dir_name:
        tmp_dir = Path(tmp_dir_name)

        front_doc = Document(template_docx)
        replace_text_in_doc(front_doc, replacements)
        remove_unwanted_front_sections(front_doc)
        trim_doc_from_text(front_doc, "摘　要")
        front_path = save_temp_doc(front_doc, tmp_dir / "front.docx")

        main_doc = Document(main_docx)
        transform_main_doc(main_doc)
        main_path = save_temp_doc(main_doc, tmp_dir / "main.docx")

        back_doc = Document(backmatter_docx)
        transform_backmatter_doc(back_doc)
        back_path = save_temp_doc(back_doc, tmp_dir / "back.docx")

        master = Document(front_path)
        composer = Composer(master)
        composer.append(Document(main_path))
        composer.append(Document(back_path))
        composer.save(str(output_docx))

    print(f"Wrote final DOCX: {output_docx}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
