from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


LATIN_FONT = "Times New Roman"
CHINESE_FONT = "SimSun"


def ensure_rfonts(style) -> OxmlElement:
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    return rfonts


def set_style_fonts(style, latin_font: str, east_asia_font: str, size_pt: float, bold: bool | None = None) -> None:
    style.font.name = latin_font
    style.font.size = Pt(size_pt)
    if bold is not None:
        style.font.bold = bold

    rfonts = ensure_rfonts(style)
    rfonts.set(qn("w:ascii"), latin_font)
    rfonts.set(qn("w:hAnsi"), latin_font)
    rfonts.set(qn("w:cs"), latin_font)
    rfonts.set(qn("w:eastAsia"), east_asia_font)


def set_run_fonts(run, latin_font: str, east_asia_font: str, size_pt: float | None = None, bold: bool | None = None) -> None:
    run.font.name = latin_font
    if size_pt is not None:
        run.font.size = Pt(size_pt)
    if bold is not None:
        run.font.bold = bold
    rfonts = run._element.get_or_add_rPr().rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        run._element.get_or_add_rPr().append(rfonts)
    rfonts.set(qn("w:ascii"), latin_font)
    rfonts.set(qn("w:hAnsi"), latin_font)
    rfonts.set(qn("w:cs"), latin_font)
    rfonts.set(qn("w:eastAsia"), east_asia_font)


def configure_styles(doc: Document) -> None:
    styles = doc.styles

    if "Normal" in styles:
        style = styles["Normal"]
        set_style_fonts(style, LATIN_FONT, CHINESE_FONT, 12)
        style.paragraph_format.line_spacing = 1.5
        style.paragraph_format.first_line_indent = Pt(24)
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.space_after = Pt(0)
        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    if "Body Text" in styles:
        style = styles["Body Text"]
        set_style_fonts(style, LATIN_FONT, CHINESE_FONT, 12)
        style.paragraph_format.line_spacing = 1.5
        style.paragraph_format.first_line_indent = Pt(24)
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.space_after = Pt(0)
        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    if "First Paragraph" in styles:
        style = styles["First Paragraph"]
        set_style_fonts(style, LATIN_FONT, CHINESE_FONT, 12)
        style.paragraph_format.line_spacing = 1.5
        style.paragraph_format.first_line_indent = Pt(24)
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.space_after = Pt(0)
        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    if "Bibliography" in styles:
        style = styles["Bibliography"]
        set_style_fonts(style, LATIN_FONT, CHINESE_FONT, 10.5)
        style.paragraph_format.line_spacing = 1.25
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.space_after = Pt(0)
        style.paragraph_format.first_line_indent = None
        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    if "Heading 1" in styles:
        style = styles["Heading 1"]
        set_style_fonts(style, LATIN_FONT, CHINESE_FONT, 22, bold=True)
        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(12)
        style.paragraph_format.first_line_indent = None

    if "Heading 2" in styles:
        style = styles["Heading 2"]
        set_style_fonts(style, LATIN_FONT, CHINESE_FONT, 16, bold=True)
        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.first_line_indent = None

    if "Heading 3" in styles:
        style = styles["Heading 3"]
        set_style_fonts(style, LATIN_FONT, CHINESE_FONT, 14, bold=True)
        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        style.paragraph_format.space_before = Pt(6)
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.first_line_indent = None

    if "Image Caption" in styles:
        style = styles["Image Caption"]
        set_style_fonts(style, LATIN_FONT, CHINESE_FONT, 10.5)
        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        style.paragraph_format.line_spacing = 1.0
        style.paragraph_format.space_before = Pt(3)
        style.paragraph_format.space_after = Pt(3)
        style.paragraph_format.first_line_indent = None


def format_cover(doc: Document) -> None:
    nonempty = [p for p in doc.paragraphs if p.text.strip()]
    if len(nonempty) < 9:
        return

    title_para = nonempty[0]
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title_para.runs:
        set_run_fonts(run, LATIN_FONT, CHINESE_FONT, 22, True)

    subtitle_para = nonempty[1]
    subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in subtitle_para.runs:
        set_run_fonts(run, LATIN_FONT, CHINESE_FONT, 14, False)

    for para in nonempty[2:9]:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in para.runs:
            set_run_fonts(run, LATIN_FONT, CHINESE_FONT, 12, False)


def touch_paragraph_runs(doc: Document) -> None:
    target_styles = {
        "Body Text": (12, False),
        "First Paragraph": (12, False),
        "Bibliography": (10.5, False),
        "Heading 1": (22, True),
        "Heading 2": (16, True),
        "Heading 3": (14, True),
        "Image Caption": (10.5, False),
    }
    for para in doc.paragraphs:
        if para.style.name not in target_styles:
            continue
        size_pt, bold = target_styles[para.style.name]
        for run in para.runs:
            set_run_fonts(run, LATIN_FONT, CHINESE_FONT, size_pt, bold)


def main() -> int:
    parser = argparse.ArgumentParser(description="Apply consistent thesis-like styles to a DOCX file.")
    parser.add_argument("input_docx", help="Input DOCX path")
    parser.add_argument("output_docx", help="Output DOCX path")
    args = parser.parse_args()

    input_path = Path(args.input_docx).resolve()
    output_path = Path(args.output_docx).resolve()
    output_path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document(str(input_path))
    configure_styles(doc)
    format_cover(doc)
    touch_paragraph_runs(doc)
    doc.save(str(output_path))
    print(output_path)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
